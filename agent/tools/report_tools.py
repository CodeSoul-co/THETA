"""
Report Generation Tool for LangChain Agent
"""

import json
from pathlib import Path
from datetime import datetime
from langchain.tools import BaseTool
from pydantic import Field


class ReportGeneratorTool(BaseTool):
    """报告生成工具 - 生成Word分析报告"""
    
    name: str = "report_generator"
    description: str = """用于生成Word格式的分析报告。
    输入: job_id (任务ID)
    前置条件: 可视化必须已完成
    功能: 汇总分析结果，生成结构化的Word报告
    输出: report.docx文件
    使用场景: 在可视化完成后，生成最终报告供用户下载"""
    
    base_dir: Path = Field(default_factory=lambda: Path("."))
    
    def _run(self, job_id: str) -> str:
        """执行报告生成"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import yaml
            
            # 加载数据
            result_dir = self.base_dir / "result" / job_id
            topics_path = self.base_dir / "ETM" / "outputs" / "topic_words" / f"{job_id}_topics.json"
            viz_dir = self.base_dir / "visualization" / "outputs" / job_id
            
            if not topics_path.exists():
                return "错误: 主题数据不存在，请先运行etm_trainer"
            
            with open(topics_path, 'r', encoding='utf-8') as f:
                topics = json.load(f)
            
            # 创建文档
            doc = Document()
            
            # 封面
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title.add_run("\n\n\n主题模型分析报告")
            run.bold = True
            run.font.size = Pt(28)
            
            info = doc.add_paragraph()
            info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            info.add_run(f"\n\n任务ID: {job_id}")
            info.add_run(f"\n生成时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
            
            doc.add_page_break()
            
            # 第一章
            doc.add_heading('第一章 数据概览', level=1)
            doc.add_paragraph(f"本次分析共识别出 {len(topics)} 个主题。")
            
            # 第二章
            doc.add_heading('第二章 主题识别结果', level=1)
            
            for topic in topics[:10]:
                doc.add_heading(f"主题 {topic['id']}: {topic['name']}", level=2)
                keywords = ', '.join(topic.get('keywords', [])[:5])
                doc.add_paragraph(f"关键词: {keywords}")
                doc.add_paragraph(f"占比: {topic.get('proportion', 0)*100:.2f}%")
                
                # 添加词云图
                wc_path = viz_dir / f"wordcloud_topic_{topic['id']}.png"
                if wc_path.exists():
                    try:
                        doc.add_picture(str(wc_path), width=Inches(4))
                    except:
                        pass
            
            # 第三章
            doc.add_page_break()
            doc.add_heading('第三章 可视化分析', level=1)
            
            charts = [
                ("topic_distribution.png", "主题分布图"),
                ("heatmap_doc_topic.png", "文档-主题热力图"),
                ("topic_similarity.png", "主题相似度矩阵")
            ]
            
            for filename, title in charts:
                chart_path = viz_dir / filename
                if chart_path.exists():
                    doc.add_heading(title, level=2)
                    try:
                        doc.add_picture(str(chart_path), width=Inches(5))
                    except:
                        pass
            
            # 保存
            result_dir.mkdir(parents=True, exist_ok=True)
            report_path = result_dir / "report.docx"
            doc.save(str(report_path))
            
            return f"报告生成成功: {report_path}"
            
        except Exception as e:
            return f"报告生成失败: {str(e)}"
    
    async def _arun(self, job_id: str) -> str:
        return self._run(job_id)
