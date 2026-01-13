"""
Report Agent: Word报告生成Agent
基于LangChain框架，自动生成结构化Word分析报告
"""

import json
import yaml
import logging
from pathlib import Path
from typing import Dict, Any, Optional, List
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


class ReportAgent:
    """
    Word报告生成Agent
    
    职责：自动生成结构化Word分析报告
    
    输入：
    - result/{job_id}/config.yaml
    - result/{job_id}/metrics.json
    - ETM/outputs/topic_words/{job_id}_topics.json
    - visualization/outputs/{job_id}/*.png
    
    输出：
    - result/{job_id}/report.docx
    
    报告结构：
    - 封面：项目信息、生成时间
    - 第一章：数据概览
    - 第二章：方法说明
    - 第三章：主题识别结果（含K张词云）
    - 第四章：主题分布分析
    - 第五章：质量评估
    - 附录：完整参数配置
    """
    
    def __init__(self, base_dir: str = "/root/autodl-tmp", llm_config: Optional[Dict] = None):
        self.base_dir = Path(base_dir)
        self.llm_config = llm_config or {}
        self.logger = self._setup_logger()
    
    def _setup_logger(self) -> logging.Logger:
        """Setup logger for this agent"""
        logger = logging.getLogger(f"ReportAgent_{id(self)}")
        logger.setLevel(logging.INFO)
        return logger
    
    def process(self, job_id: str) -> Dict[str, Any]:
        """
        生成Word报告
        
        Args:
            job_id: 任务ID
            
        Returns:
            Dict with processing results and status
        """
        try:
            self.logger.info(f"Starting report generation for job_id: {job_id}")
            
            # 加载所需数据
            config_data = self._load_config(job_id)
            metrics_data = self._load_metrics(job_id)
            topics_data = self._load_topics(job_id)
            analysis_result = self._load_analysis_result(job_id)
            
            # 创建Word文档
            doc = Document()
            
            # 设置文档样式
            self._setup_styles(doc)
            
            # 生成报告各章节
            self._add_cover_page(doc, job_id, config_data)
            self._add_chapter_1_data_overview(doc, config_data, analysis_result)
            self._add_chapter_2_methodology(doc)
            self._add_chapter_3_topic_results(doc, job_id, topics_data)
            self._add_chapter_4_distribution_analysis(doc, job_id, analysis_result)
            self._add_chapter_5_quality_evaluation(doc, metrics_data)
            self._add_appendix(doc, config_data)
            
            # 保存报告
            report_path = self.base_dir / "result" / job_id / "report.docx"
            doc.save(str(report_path))
            
            self._log(job_id, f"Report generated successfully: {report_path}")
            
            return {
                "status": "success",
                "job_id": job_id,
                "report_path": str(report_path),
                "chapters": [
                    "封面",
                    "第一章：数据概览",
                    "第二章：方法说明",
                    "第三章：主题识别结果",
                    "第四章：主题分布分析",
                    "第五章：质量评估",
                    "附录：完整参数配置"
                ]
            }
            
        except Exception as e:
            self._log(job_id, f"Report generation failed: {str(e)}", error=True)
            return {
                "status": "failed",
                "job_id": job_id,
                "error": str(e)
            }
    
    def _load_config(self, job_id: str) -> Dict:
        """加载配置文件"""
        config_path = self.base_dir / "result" / job_id / "config.yaml"
        if config_path.exists():
            try:
                with open(config_path, 'r', encoding='utf-8') as f:
                    # 使用unsafe_load处理python对象标签，或者读取为纯文本
                    content = f.read()
                    # 移除python特殊标签
                    content = content.replace('!!python/tuple', '')
                    return yaml.safe_load(content) or {}
            except Exception as e:
                self.logger.warning(f"Failed to load config: {e}")
                return {}
        return {}
    
    def _load_metrics(self, job_id: str) -> Dict:
        """加载评估指标"""
        metrics_path = self.base_dir / "result" / job_id / "metrics.json"
        if metrics_path.exists():
            with open(metrics_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {}
    
    def _load_topics(self, job_id: str) -> List[Dict]:
        """加载主题数据"""
        topics_path = self.base_dir / "ETM" / "outputs" / "topic_words" / f"{job_id}_topics.json"
        if topics_path.exists():
            with open(topics_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        return []
    
    def _load_analysis_result(self, job_id: str) -> Dict:
        """加载分析结果"""
        result_path = self.base_dir / "result" / job_id / "analysis_result.json"
        if result_path.exists():
            with open(result_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {}
    
    def _setup_styles(self, doc: Document):
        """设置文档样式"""
        # 标题样式已内置，这里可以自定义
        pass
    
    def _add_cover_page(self, doc: Document, job_id: str, config_data: Dict):
        """添加封面"""
        # 标题
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("\n\n\n主题模型分析报告")
        run.bold = True
        run.font.size = Pt(28)
        
        # 副标题
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"\nTopic Model Analysis Report")
        run.font.size = Pt(18)
        run.italic = True
        
        # 项目信息
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.add_run(f"\n\n\n任务ID: {job_id}")
        
        # 生成时间
        time_para = doc.add_paragraph()
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        time_para.add_run(f"\n生成时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
        
        # 分页
        doc.add_page_break()
    
    def _add_chapter_1_data_overview(self, doc: Document, config_data: Dict, analysis_result: Dict):
        """第一章：数据概览"""
        doc.add_heading('第一章 数据概览', level=1)
        
        doc.add_heading('1.1 数据基本信息', level=2)
        
        # 从分析结果中提取信息
        job_id = analysis_result.get('job_id', 'N/A')
        status = analysis_result.get('status', 'N/A')
        completed_at = analysis_result.get('completed_at', 'N/A')
        duration = analysis_result.get('duration_seconds', 'N/A')
        
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        rows_data = [
            ('任务ID', job_id),
            ('处理状态', status),
            ('完成时间', completed_at),
            ('处理耗时', f"{duration}秒" if duration != 'N/A' else 'N/A'),
            ('主题数量', str(len(analysis_result.get('topics', []))))
        ]
        
        for i, (key, value) in enumerate(rows_data):
            table.rows[i].cells[0].text = key
            table.rows[i].cells[1].text = str(value)
        
        doc.add_paragraph()
        
        doc.add_heading('1.2 数据来源', level=2)
        doc.add_paragraph(
            "本报告基于用户提供的文本数据进行主题建模分析。"
            "数据经过预处理后，使用嵌入式主题模型（ETM）进行主题发现。"
        )
        
        doc.add_page_break()
    
    def _add_chapter_2_methodology(self, doc: Document):
        """第二章：方法说明"""
        doc.add_heading('第二章 方法说明', level=1)
        
        doc.add_heading('2.1 嵌入式主题模型（ETM）', level=2)
        doc.add_paragraph(
            "嵌入式主题模型（Embedded Topic Model, ETM）是一种结合词嵌入和传统主题模型的方法。"
            "与传统LDA不同，ETM利用预训练的词向量来捕捉词汇的语义信息，从而生成更具语义一致性的主题。"
        )
        
        doc.add_heading('2.2 处理流程', level=2)
        
        steps = [
            "1. 数据清洗：验证输入数据格式，生成配置文件",
            "2. 词袋生成：构建词汇表和词袋矩阵（BOW）",
            "3. 嵌入生成：生成文档嵌入和词汇嵌入",
            "4. 主题建模：运行ETM模型，生成主题分布",
            "5. 可视化：生成词云、热力图等可视化结果",
            "6. 报告生成：汇总分析结果，生成结构化报告"
        ]
        
        for step in steps:
            doc.add_paragraph(step, style='List Bullet')
        
        doc.add_heading('2.3 评估指标', level=2)
        doc.add_paragraph(
            "主题一致性（Coherence Score）：衡量主题内词汇的语义一致性，分数越高表示主题质量越好。\n"
            "主题多样性（Diversity Score）：衡量不同主题之间的差异程度，分数越高表示主题越独特。"
        )
        
        doc.add_page_break()
    
    def _add_chapter_3_topic_results(self, doc: Document, job_id: str, topics_data: List[Dict]):
        """第三章：主题识别结果"""
        doc.add_heading('第三章 主题识别结果', level=1)
        
        doc.add_heading('3.1 主题概览', level=2)
        doc.add_paragraph(f"本次分析共识别出 {len(topics_data)} 个主题。以下是各主题的关键词和占比：")
        
        # 主题表格
        if topics_data:
            table = doc.add_table(rows=len(topics_data) + 1, cols=4)
            table.style = 'Table Grid'
            
            # 表头
            headers = ['主题ID', '主题名称', '关键词', '占比']
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = header
            
            # 数据行
            for i, topic in enumerate(topics_data):
                table.rows[i + 1].cells[0].text = str(topic.get('id', i))
                table.rows[i + 1].cells[1].text = topic.get('name', f'Topic_{i}')
                keywords = topic.get('keywords', [])[:5]  # 只显示前5个关键词
                table.rows[i + 1].cells[2].text = ', '.join(keywords)
                proportion = topic.get('proportion', 0)
                table.rows[i + 1].cells[3].text = f"{proportion*100:.2f}%"
        
        doc.add_paragraph()
        
        doc.add_heading('3.2 主题词云', level=2)
        doc.add_paragraph("以下是各主题的词云可视化，词语大小表示其在该主题中的重要性：")
        
        # 添加词云图片（最多显示10个）
        viz_dir = self.base_dir / "visualization" / "outputs" / job_id
        for i in range(min(10, len(topics_data))):
            wordcloud_path = viz_dir / f"wordcloud_topic_{i}.png"
            if wordcloud_path.exists():
                doc.add_paragraph(f"\n主题 {i} 词云：")
                try:
                    doc.add_picture(str(wordcloud_path), width=Inches(5))
                except Exception as e:
                    doc.add_paragraph(f"[词云图片加载失败: {e}]")
        
        doc.add_page_break()
    
    def _add_chapter_4_distribution_analysis(self, doc: Document, job_id: str, analysis_result: Dict):
        """第四章：主题分布分析"""
        doc.add_heading('第四章 主题分布分析', level=1)
        
        doc.add_heading('4.1 主题分布图', level=2)
        doc.add_paragraph("以下图表展示了各主题在文档集中的分布情况：")
        
        # 添加主题分布图
        dist_path = self.base_dir / "visualization" / "outputs" / job_id / "topic_distribution.png"
        if dist_path.exists():
            try:
                doc.add_picture(str(dist_path), width=Inches(6))
            except Exception as e:
                doc.add_paragraph(f"[图片加载失败: {e}]")
        
        doc.add_heading('4.2 文档-主题热力图', level=2)
        doc.add_paragraph("热力图展示了每个文档与各主题的关联强度，颜色越深表示关联越强：")
        
        heatmap_path = self.base_dir / "visualization" / "outputs" / job_id / "heatmap_doc_topic.png"
        if heatmap_path.exists():
            try:
                doc.add_picture(str(heatmap_path), width=Inches(6))
            except Exception as e:
                doc.add_paragraph(f"[图片加载失败: {e}]")
        
        doc.add_heading('4.3 主题相似度矩阵', level=2)
        doc.add_paragraph("相似度矩阵展示了主题之间的语义相似程度：")
        
        similarity_path = self.base_dir / "visualization" / "outputs" / job_id / "topic_similarity.png"
        if similarity_path.exists():
            try:
                doc.add_picture(str(similarity_path), width=Inches(5))
            except Exception as e:
                doc.add_paragraph(f"[图片加载失败: {e}]")
        
        doc.add_page_break()
    
    def _add_chapter_5_quality_evaluation(self, doc: Document, metrics_data: Dict):
        """第五章：质量评估"""
        doc.add_heading('第五章 质量评估', level=1)
        
        doc.add_heading('5.1 评估指标', level=2)
        
        coherence = metrics_data.get('coherence_score', 'N/A')
        diversity = metrics_data.get('diversity_score', 'N/A')
        optimal_k = metrics_data.get('optimal_k', 'N/A')
        
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Table Grid'
        
        # 表头
        table.rows[0].cells[0].text = '指标'
        table.rows[0].cells[1].text = '数值'
        table.rows[0].cells[2].text = '解读'
        
        # 数据
        table.rows[1].cells[0].text = '主题一致性'
        table.rows[1].cells[1].text = str(coherence)
        table.rows[1].cells[2].text = '> 0.7 表示主题质量良好' if coherence != 'N/A' else ''
        
        table.rows[2].cells[0].text = '主题多样性'
        table.rows[2].cells[1].text = str(diversity)
        table.rows[2].cells[2].text = '> 0.5 表示主题差异明显' if diversity != 'N/A' else ''
        
        table.rows[3].cells[0].text = '最优主题数'
        table.rows[3].cells[1].text = str(optimal_k)
        table.rows[3].cells[2].text = '模型选择的最佳主题数量'
        
        doc.add_paragraph()
        
        doc.add_heading('5.2 质量评价', level=2)
        
        if coherence != 'N/A' and float(coherence) >= 0.7:
            doc.add_paragraph("✓ 主题一致性良好：模型生成的主题内部语义连贯，关键词之间具有较强的语义关联。")
        else:
            doc.add_paragraph("△ 主题一致性一般：建议调整主题数量或优化数据预处理。")
        
        if diversity != 'N/A' and float(diversity) >= 0.5:
            doc.add_paragraph("✓ 主题多样性良好：各主题之间差异明显，没有明显的主题重叠。")
        else:
            doc.add_paragraph("△ 主题多样性一般：部分主题可能存在重叠，建议减少主题数量。")
        
        doc.add_page_break()
    
    def _add_appendix(self, doc: Document, config_data: Dict):
        """附录：完整参数配置"""
        doc.add_heading('附录 完整参数配置', level=1)
        
        doc.add_paragraph("以下是本次分析使用的完整参数配置：")
        
        if config_data:
            # 将配置转为格式化文本
            config_text = yaml.dump(config_data, allow_unicode=True, default_flow_style=False)
            doc.add_paragraph(config_text, style='No Spacing')
        else:
            doc.add_paragraph("配置文件未找到或为空。")
        
        doc.add_paragraph("\n\n--- 报告结束 ---")
    
    def _log(self, job_id: str, message: str, error: bool = False):
        """Log message to result/{job_id}/log.txt"""
        log_path = self.base_dir / "result" / job_id / "log.txt"
        log_path.parent.mkdir(parents=True, exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        level = "ERROR" if error else "INFO"
        log_entry = f"[{timestamp}] [{level}] ReportAgent: {message}\n"
        
        with open(log_path, 'a', encoding='utf-8') as f:
            f.write(log_entry)
        
        if error:
            self.logger.error(message)
        else:
            self.logger.info(message)
